"""
组卷模块 — CET-4 智能组卷 + 试卷管理 + 在线考试
按 CET4 固定 57 题结构组卷，对接 difficulty 模块按难度比例抽题。
彻底修复"组卷后题目显示混乱"Bug：内存显式排序 + Passage 块级聚合。
"""
import random
import json
import re
import os
import difflib
from datetime import datetime
from collections import OrderedDict, defaultdict
from sqlalchemy import func
from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel
from db.database import SessionLocal
from models.question import Question as DbQuestion
from models.paper import Paper, ExamPaperMapping
from models.student_answer import StudentAnswer
from models.exam_record import ExamRecord
from utils.auth import require_teacher

router = APIRouter(tags=["组卷与考试"])

# ── CET4 固定四级标准卷面结构 ──────────────────────────────────────────
# 严格规范 question_id 1-57 的卷面顺延顺序
# 题号 1：    写作（Part I Writing）
# 题号 2-26： 听力（Part II Listening，共 25 道子题）
# 题号 27-56：阅读（Part III Reading，共 30 道子题）
#             27-36 选词填空  37-46 段落匹配
#             47-51 仔细阅读 Passage One  52-56 仔细阅读 Passage Two
# 题号 57：   翻译（Part IV Translation）
CET4_LAYOUT = [
    (1, "写作", "Part I Writing", 106.5),
]
# 听力 Section A-B: qid 2-16 (15题, 每题7.1分)
for qid in range(2, 17):
    CET4_LAYOUT.append((qid, "听力", "Part II Listening (Sections A-B)", 7.1))
# 听力 Section C: qid 17-26 (10题, 每题14.2分)
for qid in range(17, 27):
    CET4_LAYOUT.append((qid, "听力", "Part II Listening (Section C)", 14.2))
# 选词填空: qid 27-36 (10题, 每题3.55分)
for qid in range(27, 37):
    CET4_LAYOUT.append((qid, "选词填空", "Part III Reading (Section A)", 3.55))
# 段落匹配: qid 37-46 (10题, 每题7.1分)
for qid in range(37, 47):
    CET4_LAYOUT.append((qid, "段落匹配", "Part III Reading (Section B)", 7.1))
# 仔细阅读: qid 47-56 (10题, 每题14.2分)
for qid in range(47, 57):
    CET4_LAYOUT.append((qid, "仔细阅读", "Part III Reading (Section C)", 14.2))
# 翻译: qid 57 (1题, 106.5分)
CET4_LAYOUT.append((57, "翻译", "Part IV Translation", 106.5))

# 题型 → 数据库 type 字段值映射（兼容中英文）
TYPE_MAP = {
    "写作":     ["写作", "essay", "writing"],
    "听力":     ["听力", "listening"],
    "选词填空": ["选词填空", "fill", "cloze"],
    "段落匹配": ["段落匹配", "matching"],
    "仔细阅读": ["仔细阅读", "reading", "阅读"],
    "翻译":     ["翻译", "translation"],
}

# ── 难度权重（对接 difficulty 模块 1-5 级）────────────────────────
DIFF_WEIGHTS = {
    1: {1: 0.45, 2: 0.30, 3: 0.15, 4: 0.07, 5: 0.03},
    2: {1: 0.25, 2: 0.35, 3: 0.25, 4: 0.10, 5: 0.05},
    3: {1: 0.10, 2: 0.20, 3: 0.40, 4: 0.20, 5: 0.10},
    4: {1: 0.05, 2: 0.10, 3: 0.25, 4: 0.35, 5: 0.25},
    5: {1: 0.03, 2: 0.07, 3: 0.15, 4: 0.30, 5: 0.45},
}
DIFF_LABEL = {1: "基础", 2: "进阶", 3: "中等", 4: "较难", 5: "困难"}

# ── 组卷参数 ───────────────────────────────────────────────────────────

class PaperGenerateReq(BaseModel):
    difficulty: int = 3
    title: str = ""

# ── 题干清洗 ───────────────────────────────────────────────────────────

def _clean_stem_content(content):
    """清理题干杂质，彻底移除原文段落"""
    if not content:
        return content
    
    # ★ 优化：彻底移除 [原文] 后面的所有原文内容
    if '[原文]' in content:
        # 处理 [原文]...[题目] 格式
        if '[题目]' in content:
            # 移除 [原文] 到 [题目] 之间的所有内容
            content = re.sub(r'\[原文\][\s\S]*?\[题目\]\s*', '', content)
        else:
            # 没有 [题目] 标记，说明 content 只有原文，没有题目句子
            # 对于选词填空、段落匹配等题型，直接返回空字符串
            # 因为原文应该在 passage_text 字段中，不应该在 content 中显示
            return ""
    
    # 移除 Questions X to Y are based on... 这类提示语
    content = re.sub(
        r'\s*Questions?\s*\d+\s*(?:to|and)\s*\d+\s*are\s*based\s*(?:on|upon)\s*(?:the\s+(?:following\s+)?passage|what\s+you\s+have\s+just\s+heard)\.?\s*',
        '', content, flags=re.IGNORECASE)
    # 清理多余空白
    content = re.sub(r'\s+', ' ', content).strip()
    return content


def _reconstruct_question(q):
    """重新构建题目，分离题干和选项"""
    content = q.content or ""
    options = []
    if q.options:
        try:
            options = json.loads(q.options) if isinstance(q.options, str) else q.options
        except Exception:
            options = []
    # 选项字符串列表 → 字典格式
    if options and isinstance(options[0], str):
        processed = []
        for i, opt in enumerate(options):
            if len(opt) >= 3 and opt[1] in ('.', ')'):
                processed.append({"label": opt[0], "text": opt[2:].strip()})
            else:
                m = re.match(r'^([A-D])[\.\)]\s*(.*)', opt)
                if m:
                    processed.append({"label": m.group(1), "text": m.group(2).strip()})
                else:
                    processed.append({"label": chr(65 + i), "text": opt.strip()})
        options = processed
    # 检查题干是否以选项开头
    opt_start = re.match(r'^[A-D][)\.]\s*', content)
    if opt_start:
        first_label = opt_start.group(0)[0]
        found = any(opt.get('label') == first_label for opt in options)
        if not found:
            next_opt = re.search(r'\b[B-D][)\.]\s', content)
            if next_opt:
                opt_text = content[opt_start.end():next_opt.start()]
                options.insert(0, {"label": first_label, "text": _clean_stem_content(opt_text)})
                content = content[next_opt.start():]
            else:
                end_pos = len(content)
                for lbl in ['B.', 'B)', 'C.', 'C)', 'D.', 'D)']:
                    idx = content.find(lbl)
                    if idx > 0:
                        end_pos = min(end_pos, idx)
                if end_pos > opt_start.end():
                    opt_text = content[opt_start.end():end_pos]
                    options.insert(0, {"label": first_label, "text": _clean_stem_content(opt_text)})
                    content = content[end_pos:]
    content = _clean_stem_content(content)
    if not content or len(content.strip()) < 5:
        # ★ 优化：为不同题型生成更合适的默认描述
        qtype = getattr(q, 'type', '')
        qsection = getattr(q, 'section', '')
        # 使用 section 字段判断题型
        if 'Section A' in qsection or '选词填空' in qtype:
            # 选词填空：显示填空提示
            content = "请根据原文选择正确的词语填空："
        elif 'Section B' in qsection or '段落匹配' in qtype:
            # 段落匹配：显示匹配提示
            content = "请根据原文匹配正确的段落："
        elif options:
            content = "请根据以下选项回答问题："
        else:
            content = "[题目内容]"
    options.sort(key=lambda x: x.get('label', 'Z'))
    return content, options


# ── 抽题辅助（对接 difficulty 模块 1-5 级难度）─────────────────────

def _to_difficulty_int(q):
    """将 question.difficulty 统一转为 int 1-5。兼容 String/Int 存储差异。"""
    try:
        d = int(q.difficulty or 3)
    except (ValueError, TypeError):
        d = 3
    return max(1, min(5, d))


def _pick_by_ratio(pool, count, difficulty):
    """按难度比例抽题。对接 difficulty 模块的 1-5 级难度体系。
    difficulty 参数: 1=偏基础, 2=偏进阶, 3=中等, 4=偏难, 5=偏困难
    """
    weights = DIFF_WEIGHTS.get(difficulty, DIFF_WEIGHTS[3])
    buckets = {}
    for d in range(1, 6):
        buckets[d] = [q for q in pool if _to_difficulty_int(q) == d]

    selected = []
    seen_contents = []  # 用于去重
    for d, bucket in buckets.items():
        if bucket:
            need = int(count * weights[d])
            # 从该难度桶中随机选择，但要去重
            shuffled = bucket.copy()
            random.shuffle(shuffled)
            for q in shuffled:
                if len(selected) >= count:
                    break
                if len([s for s in selected if s.id == q.id]) > 0:
                    continue
                # 检查内容相似度
                content = (getattr(q, 'content', '') or '').strip()
                is_duplicate = False
                for seen_content in seen_contents:
                    similarity = difflib.SequenceMatcher(None, content.lower(), seen_content.lower()).ratio()
                    if similarity > 0.85:
                        is_duplicate = True
                        break
                if not is_duplicate:
                    selected.append(q)
                    seen_contents.append(content)

    # 补齐不足
    if len(selected) < count:
        # 使用 id 去重
        selected_ids = {q.id for q in selected}
        # 收集所有可用题目（排除已选的）
        remaining = []
        for q in pool:
            if q.id not in selected_ids:
                content = (getattr(q, 'content', '') or '').strip()
                is_duplicate = False
                for seen_content in seen_contents:
                    similarity = difflib.SequenceMatcher(None, content.lower(), seen_content.lower()).ratio()
                    if similarity > 0.85:
                        is_duplicate = True
                        break
                if not is_duplicate:
                    remaining.append(q)
        
        if remaining:
            need = min(count - len(selected), len(remaining))
            selected.extend(random.sample(remaining, need))
    
    return selected[:count]


def _find_pool(all_qs, display_type):
    """根据展示题型名在题库中查找匹配的题目池"""
    # ★ 优化：阅读类题目按 section 字段区分
    if display_type == "选词填空":
        # Reading Section A = 选词填空
        pool = [q for q in all_qs if (q.section or '').find('Reading Section A') >= 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if (q.section or '').find('Section A') >= 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if q.type in ("阅读", "reading") and q.answer and q.answer.strip() and q.content and q.content.strip()]
        return pool
    
    if display_type == "段落匹配":
        # Reading Section B = 段落匹配
        pool = [q for q in all_qs if (q.section or '').find('Reading Section B') >= 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if (q.section or '').find('Section B') >= 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if q.type in ("阅读", "reading") and q.answer and q.answer.strip() and q.content and q.content.strip()]
        return pool
    
    if display_type == "仔细阅读":
        # Reading Section C = 仔细阅读（注意排除 Listening Section C）
        pool = [q for q in all_qs if (q.section or '').find('Reading Section C') >= 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            # 更宽松的匹配：包含Section C但不包含Listening
            pool = [q for q in all_qs if (q.section or '').find('Section C') >= 0 and (q.section or '').find('Listening') < 0 and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if q.type in ("阅读", "reading") and q.answer and q.answer.strip() and q.content and q.content.strip()]
        return pool
    
    # 写作和翻译：不检查 answer（因为它们通常没有标准答案）
    if display_type == "写作":
        pool = [q for q in all_qs if q.type in ("写作", "essay", "writing") and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if (q.section or '').find('Writing') >= 0 and q.content and q.content.strip()]
        return pool
    
    if display_type == "翻译":
        pool = [q for q in all_qs if q.type in ("翻译", "translation") and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if (q.section or '').find('Translation') >= 0 and q.content and q.content.strip()]
        return pool
    
    # 其他题型（如听力）按 type 字段匹配
    candidates = TYPE_MAP.get(display_type, [display_type])
    pool = []
    for cand in candidates:
        pool.extend([q for q in all_qs if q.type == cand and q.answer and q.answer.strip() and q.content and q.content.strip()])
    
    if not pool:
        if display_type in ("选词填空", "段落匹配", "仔细阅读"):
            pool = [q for q in all_qs if q.type in ("阅读", "reading") and q.answer and q.answer.strip() and q.content and q.content.strip()]
        if not pool:
            pool = [q for q in all_qs if q.answer and q.answer.strip() and q.content and q.content.strip()]
    
    # 去重
    seen = set()
    unique = []
    for q in pool:
        if q.id not in seen:
            seen.add(q.id)
            unique.append(q)
    return unique


# ── Passage 分组（拒绝扁平化，构建嵌套树状结构）─────────────────────

def _build_passage_groups(pool, group_key='passage_text'):
    """将题目按原文 passage 或 section 分组。
    返回 [{passage: str, passage_key: str, questions: [q,...]}, ...]
    每组内按 question_number 或 id 排序。
    """
    groups = defaultdict(list)
    for q in pool:
        if group_key == 'passage_text':
            key = (getattr(q, 'passage_text', '') or '').strip()
            if not key:
                key = f'__no_passage_{q.id}__'
        elif group_key == 'section':
            # ★ 新增：按所属大题（section）分组，保证题目连贯性
            key = (getattr(q, 'section', '') or '').strip()
            if not key:
                key = f'__no_section_{q.id}__'
        else:
            key = (getattr(q, 'source', '') or '').strip()
            if not key:
                key = f'__no_source_{q.id}__'
        groups[key].append(q)

    result = []
    for key, qs in groups.items():
        # 组内按 question_number 排序；无 question_number 按 id
        qs_sorted = sorted(qs, key=lambda x: (getattr(x, 'question_number', 0) or 0, x.id))
        # 取第一个有 passage_text 的作为 passage
        passage = ''
        for q in qs_sorted:
            pt = (getattr(q, 'passage_text', '') or '').strip()
            if pt:
                passage = pt
                break
        if group_key not in ('passage_text', 'section'):
            passage = ''
        result.append({'passage': passage, 'passage_key': key, 'questions': qs_sorted})

    # 按组大小降序
    result.sort(key=lambda g: -len(g['questions']))
    return result


def _pick_from_groups(groups, total_needed, difficulty, pick_one_group=False):
    """从 passage 组中按需抽取题目，尽量保持 passage 完整。
    pick_one_group=True: 从单个组取（选词填空保持同源）
    """
    selected = []
    if pick_one_group:
        for g in groups:
            if len(g['questions']) >= total_needed:
                selected = _pick_by_ratio(g['questions'], total_needed, difficulty)
                break
        if not selected:
            best = max(groups, key=lambda g: len(g['questions'])) if groups else None
            if best:
                selected = _pick_by_ratio(best['questions'], total_needed, difficulty)
        
        # 基于内容去重（防止数据库中存在重复或高度相似题目）
        final_selected = []
        seen_contents = []  # 存储已选题目内容用于相似度比较
        for q in selected:
            content = (getattr(q, 'content', '') or '').strip()
            if not content:
                continue
            
            # 检查是否与已选题目高度相似（相似度 > 0.85）
            is_duplicate = False
            for seen_content in seen_contents:
                similarity = difflib.SequenceMatcher(None, content.lower(), seen_content.lower()).ratio()
                if similarity > 0.85:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                seen_contents.append(content)
                final_selected.append(q)
        
        # 如果去重后数量不够，补充题目
        if len(final_selected) < total_needed and groups:
            best_group = max(groups, key=lambda g: len(g['questions']))
            # 收集所有可用题目（排除已选的）
            available = []
            selected_ids = {q.id for q in final_selected}
            for q in best_group['questions']:
                if q.id not in selected_ids:
                    content_key = (getattr(q, 'content', '') or '').strip()
                    normalized_key = ''.join(content_key.lower().split())[:100]
                    if normalized_key and normalized_key not in seen_content:
                        available.append(q)
            
            # 随机补充
            if available:
                need = total_needed - len(final_selected)
                final_selected.extend(random.sample(available, min(need, len(available))))
        
        return final_selected[:total_needed]

    # 常规：尽量取完整组
    remaining = total_needed
    for g in groups:
        if remaining <= 0:
            break
        if len(g['questions']) <= remaining:
            selected.extend(g['questions'])
            remaining -= len(g['questions'])

    if remaining > 0:
        already_ids = {q.id for q in selected}
        for g in groups:
            if remaining <= 0:
                break
            leftover = [q for q in g['questions'] if q.id not in already_ids]
            if leftover:
                taken = _pick_by_ratio(leftover, min(remaining, len(leftover)), difficulty)
                selected.extend(taken)
                remaining -= len(taken)
    
    # 基于内容去重（防止数据库中存在重复或高度相似题目）
    final_selected = []
    seen_contents = []  # 存储已选题目内容用于相似度比较
    for q in selected:
        content = (getattr(q, 'content', '') or '').strip()
        if not content:
            continue
        
        # 检查是否与已选题目高度相似（相似度 > 0.9）
        is_duplicate = False
        for seen_content in seen_contents:
            similarity = difflib.SequenceMatcher(None, content.lower(), seen_content.lower()).ratio()
            if similarity > 0.9:
                is_duplicate = True
                break
        
        if not is_duplicate:
            seen_contents.append(content)
            final_selected.append(q)
    
    # 如果去重后数量不够，补充题目
    if len(final_selected) < total_needed and groups:
        # 收集所有可用题目（排除已选的）
        available = []
        selected_ids = {q.id for q in final_selected}
        for g in groups:
            for q in g['questions']:
                if q.id not in selected_ids:
                    content_key = (getattr(q, 'content', '') or '').strip()
                    normalized_key = ''.join(content_key.lower().split())[:100]
                    if normalized_key and normalized_key not in seen_content:
                        available.append(q)
        
        # 随机补充
        if available:
            need = total_needed - len(final_selected)
            final_selected.extend(random.sample(available, min(need, len(available))))
    
    return final_selected[:total_needed]


# ══════════════════════════════════════════════════════════════════════════
# 核心：智能组卷（修复乱序 + Passage 块级聚合 + 持久化映射）
# ══════════════════════════════════════════════════════════════════════════

@router.post("/api/paper/generate")
def generate_paper(req: PaperGenerateReq, teacher = Depends(require_teacher)):
    """按 CET4 固定 57 题结构生成试卷。
    严格规范 question_id 1-57 卷面顺延顺序。
    阅读题型按 passage 分组，原文保持完整。
    """
    try:
        db = SessionLocal()
        try:
            all_qs = db.query(DbQuestion).all()
            if not all_qs:
                return {"code": 500, "msg": "题库为空，请先导入题目"}

            # ── 1. 按题型建池 ──
            pools = {}
            for display_type in ["写作", "听力", "翻译"]:
                pools[display_type] = _find_pool(all_qs, display_type)
                if not pools[display_type]:
                    pools[display_type] = all_qs

            # 阅读类分别建池
            fill_pool = _find_pool(all_qs, "选词填空")
            match_pool = _find_pool(all_qs, "段落匹配")
            read_pool = _find_pool(all_qs, "仔细阅读")

            # 去重：同一题不应出现在多个池中，但允许 fallback
            if not fill_pool:
                fill_pool = [q for q in all_qs if q.type in ("阅读", "reading")]
            if not match_pool:
                match_pool = [q for q in all_qs if q.type in ("阅读", "reading")]
            if not read_pool:
                read_pool = [q for q in all_qs if q.type in ("阅读", "reading")]

            # 统计需求量
            type_needed = {}
            for _, qtype, _, _ in CET4_LAYOUT:
                type_needed[qtype] = type_needed.get(qtype, 0) + 1

            # ── 2. 按难度比例抽题 ──
            type_selected = {}
            # 写作和翻译：独立抽题
            for qtype in ["写作", "翻译"]:
                need = type_needed.get(qtype, 0)
                if need > 0:
                    selected = _pick_by_ratio(pools[qtype], need, req.difficulty)
                    # ★ 显式排序：修复乱序 Bug
                    selected.sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))
                    type_selected[qtype] = selected
            
            # ★ 听力：按 source 分组抽题（保持同一套听力题的连贯性）
            listening_need = type_needed.get("听力", 0)
            if listening_need > 0:
                listening_groups = _build_passage_groups(pools["听力"], group_key='source')
                listening_groups = [g for g in listening_groups if len(g['questions']) >= 5]  # 每组至少5题
                type_selected["听力"] = _pick_from_groups(listening_groups, listening_need, req.difficulty)
                type_selected["听力"].sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))

            # 选词填空：按 section 分组，从同一所属大题取题，保证题目连贯性
            fill_need = type_needed.get("选词填空", 0)
            # ★ 优先按 section 分组（所属大题），保证题目完整连贯
            fill_groups = _build_passage_groups(fill_pool, group_key='section')
            fill_groups = [g for g in fill_groups if len(g['questions']) >= 5]
            if not fill_groups:
                # 如果 section 分组不行，再回退到 source 分组
                fill_groups = _build_passage_groups(fill_pool, group_key='source')
                fill_groups = [g for g in fill_groups if len(g['questions']) >= fill_need]
            type_selected["选词填空"] = _pick_from_groups(
                fill_groups, fill_need, req.difficulty, pick_one_group=True)
            # ★ 组内排序
            type_selected["选词填空"].sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))

            # 段落匹配：优先按 passage_text 分组，保证题目连贯性（每篇文章10题）
            pm_need = type_needed.get("段落匹配", 0)
            # ★ 优先按 passage_text 分组（每篇文章对应10道题）
            pm_groups = _build_passage_groups(match_pool, group_key='passage_text')
            pm_groups = [g for g in pm_groups if len(g['questions']) >= 10]
            if not pm_groups:
                # 如果 passage_text 分组不行，再尝试按 section 分组
                pm_groups = _build_passage_groups(match_pool, group_key='section')
                pm_groups = [g for g in pm_groups if len(g['questions']) >= 5]
            if not pm_groups:
                # 最后回退到 source 分组
                pm_groups = _build_passage_groups(match_pool, group_key='source')
                pm_groups = [g for g in pm_groups if len(g['questions']) >= 3]
            
            # ★ 随机打乱groups顺序，使每次选择不同的passage
            random.shuffle(pm_groups)
            
            type_selected["段落匹配"] = _pick_from_groups(pm_groups, pm_need, req.difficulty, pick_one_group=True)
            type_selected["段落匹配"].sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))
            used_ids = {q.id for q in type_selected.get("段落匹配", [])}

            # 仔细阅读：按 section 分组，保证题目连贯性
            cr_need = type_needed.get("仔细阅读", 0)
            # ★ 优先按 section 分组（所属大题）
            cr_groups = _build_passage_groups(read_pool, group_key='section')
            cr_groups = [g for g in cr_groups if len(g['questions']) >= 5]
            if not cr_groups:
                # 如果 section 分组不行，再回退到 passage_text 分组
                cr_groups = _build_passage_groups(read_pool, group_key='passage_text')
                cr_groups = [g for g in cr_groups if len(g['questions']) >= 3]
            # 移除已用于段落匹配的题目
            for g in cr_groups:
                g['questions'] = [q for q in g['questions'] if q.id not in used_ids]
            cr_groups = [g for g in cr_groups if len(g['questions']) >= 3]
            type_selected["仔细阅读"] = _pick_from_groups(cr_groups, cr_need, req.difficulty)
            type_selected["仔细阅读"].sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))

            # ── 3. 按 CET4_LAYOUT 顺序编排（核心排序逻辑）──
            # ★ 写死升序排序函数：ORDER BY question_id ASC
            type_idx = {}
            result = []  # [(qid, qtype, section, full_score, db_question), ...]
            for qid, qtype, section, full_score in CET4_LAYOUT:
                lst = type_selected.get(qtype, [])
                idx = type_idx.get(qtype, 0)
                if idx < len(lst):
                    q = lst[idx]
                    type_idx[qtype] = idx + 1
                    result.append((qid, qtype, section, full_score, q))
                else:
                    # 如果该题型题库不够，记录警告但继续
                    pass

            if not result:
                return {"code": 500, "msg": "题库中没有可用题目，请先导入题库"}

            # ── 4. 构建 CET-4 标准试卷结构（Passage 块级聚合）──
            reading_types = {"选词填空", "段落匹配", "仔细阅读"}

            # 4a. 收集各题型题目数据（含 passage 信息）
            typed_items = OrderedDict()
            for qid, qtype, section, full_score, q in result:
                content, options = _reconstruct_question(q)
                d = _to_difficulty_int(q)
                passage_text = ""
                if qtype in reading_types:
                    passage_text = (getattr(q, 'passage_text', '') or '').strip()
                    # 如果 passage_text 为空，尝试从 content 中提取 [原文] 部分
                    if not passage_text and q.content:
                        raw = q.content or ""
                        m = re.search(r'\[原文\]\s*([\s\S]*?)(?:\[题目\]|$)', raw)
                        if m:
                            passage_text = m.group(1).strip()
                    # 从题干中剥离重复原文
                    if passage_text and len(passage_text) > 50:
                        pt_head = passage_text[:80].strip()
                        ct = content.strip()
                        idx_pos = ct.find(pt_head)
                        if idx_pos >= 0:
                            after = ct[idx_pos + len(passage_text):].strip()
                            if len(after) > 5:
                                content = after
                qdata = {
                    "qid": qid,
                    "id": q.id,
                    "type": qtype,
                    "section": section,
                    "content": content,
                    "options": options,
                    "answer": q.answer or "",
                    "analysis": q.analysis or "",
                    "difficulty": d,
                    "difficulty_label": DIFF_LABEL.get(d, "未知"),
                    "full_score": full_score,
                    "passage": passage_text,
                }
                typed_items.setdefault(qtype, []).append(qdata)

            # 4b. 构建 Passage Groups（阅读题型嵌套结构）
            def _make_passage_groups(items, qtype):
                """将题目按 passage 分组。
                选词填空：单组，10题共用一个 Word Bank。
                段落匹配：单组，10题共用一个长 passage。
                仔细阅读：双 Passage，Passage One (47-51), Passage Two (52-56)。
                """
                if qtype == "选词填空":
                    qs = [{k: v for k, v in item.items() if k != "passage"} for item in items]
                    pg_passage = ""
                    word_bank = []
                    # 尝试从题干内容中提取 Word Bank 原文
                    for item in items:
                        if item.get("passage") and len(item["passage"]) > 100:
                            pg_passage = item["passage"]
                        # 收集 Word Bank 选项（从第一个有选项的题目中获取）
                        if not word_bank and item.get("options"):
                            word_bank = item["options"]
                    # 如果没有 passage_text，尝试从第一个题目的原文信息重建
                    if not pg_passage:
                        # 查找同一 source 下带有 passage_text 的题目
                        for item in items:
                            raw_content = item.get("content", "")
                            # 检查题干是否包含原文标记
                            if "[原文]" in raw_content:
                                pg_passage = raw_content
                                break
                    return [{"passage": pg_passage, "word_bank": word_bank, "questions": qs}]

                elif qtype == "仔细阅读":
                    # ★ 关键：分成两个 Passage (Passage One: qid 47-51, Passage Two: qid 52-56)
                    pg_map = OrderedDict()
                    for item in items:
                        key = item.get("passage", "") or f"__no_{item['id']}__"
                        if key not in pg_map:
                            pg_map[key] = {"passage": item.get("passage", ""), "questions": []}
                        pg_map[key]["questions"].append(
                            {k: v for k, v in item.items() if k != "passage"})

                    groups = list(pg_map.values())
                    # 按 qid 排序每个组内题目
                    for g in groups:
                        g["questions"].sort(key=lambda x: x.get("qid", 0))

                    # ★ 优化：收集所有题目，确保恰好是10道题，分成2组各5题
                    all_qs = []
                    all_passages = []
                    for g in groups:
                        all_qs.extend(g["questions"])
                        if g["passage"]:
                            all_passages.append(g["passage"])
                    
                    # 按 qid 排序所有题目
                    all_qs.sort(key=lambda x: x.get("qid", 0))
                    
                    # 确保我们有至少10道题（实际上应该有）
                    if len(all_qs) < 10:
                        # 如果不够，填充（实际上不应该发生）
                        pass
                    
                    # 取前10道题
                    all_qs = all_qs[:10]
                    
                    # 确定两个Passage的原文
                    passage1 = all_passages[0] if all_passages else ""
                    passage2 = all_passages[1] if len(all_passages) > 1 else passage1
                    
                    # 如果只有一个passage，就共用
                    if len(all_passages) == 1:
                        passage2 = passage1
                    
                    # 分成两个Passage：前5题，后5题
                    return [
                        {"passage": passage1, "questions": all_qs[:5]},
                        {"passage": passage2, "questions": all_qs[5:10]},
                    ]

                else:  # 段落匹配：单组，10题共用一个长 passage
                    pg_map = OrderedDict()
                    for item in items:
                        key = item.get("passage", "") or f"__no_{item['id']}__"
                        if key not in pg_map:
                            pg_map[key] = {"passage": item.get("passage", ""), "questions": []}
                        pg_map[key]["questions"].append(
                            {k: v for k, v in item.items() if k != "passage"})
                    for g in pg_map.values():
                        g["questions"].sort(key=lambda x: x.get("qid", 0))
                    return list(pg_map.values())

            fill_groups = _make_passage_groups(typed_items.get("选词填空", []), "选词填空")
            match_groups = _make_passage_groups(typed_items.get("段落匹配", []), "段落匹配")
            reading_groups = _make_passage_groups(typed_items.get("仔细阅读", []), "仔细阅读")

            # ── 5. CET-4 标准分层结构 ──
            paper = {
                "title": req.title or f"四级模拟试卷_{datetime.now().strftime('%Y%m%d%H%M%S')}",
                "difficulty": req.difficulty,
                "total_score": 710,
                "structure": {
                    "写作":  {"qid_range": "1",     "count": 1,  "full_score": 106.5},
                    "听力":  {"qid_range": "2-26",  "count": 25, "full_score": 248.5},
                    "选词填空": {"qid_range": "27-36", "count": 10, "full_score": 35.5},
                    "段落匹配": {"qid_range": "37-46", "count": 10, "full_score": 71},
                    "仔细阅读": {"qid_range": "47-56", "count": 10, "full_score": 142},
                    "翻译":  {"qid_range": "57",    "count": 1,  "full_score": 106.5},
                },
                "sections": [
                    {
                        "part": "Part I 写作",
                        "type": "写作",
                        "duration": "30分钟",
                        "instruction": "Directions: For this part, you are allowed 30 minutes to write an essay. You should write at least 120 words but no more than 180 words.",
                        "questions": typed_items.get("写作", []),
                    },
                    {
                        "part": "Part II 听力理解",
                        "type": "听力",
                        "duration": "25分钟",
                        "instruction": "Directions: In this section, you will hear news reports, long conversations and passages. At the end of each, you will hear some questions. After you hear a question, choose the best answer from the four choices.",
                        "questions": typed_items.get("听力", []),
                    },
                    {
                        "part": "Part III 阅读理解",
                        "type": "阅读",
                        "duration": "40分钟",
                        "instruction": "Directions: There are 3 passages in this section. Each passage is followed by some questions or unfinished statements. For each of them there are four choices marked A), B), C) and D). You should decide on the best choice.",
                        "subsections": [
                            {
                                "title": "Section A 选词填空",
                                "type": "选词填空",
                                "qid_range": "27-36",
                                "instruction": "Directions: In this section, there is a passage with ten blanks. You are required to select one word for each blank from a list of choices given in a word bank following the passage.",
                                "groups": fill_groups,
                            },
                            {
                                "title": "Section B 段落匹配",
                                "type": "段落匹配",
                                "qid_range": "37-46",
                                "instruction": "Directions: In this section, you are going to read a passage with ten statements attached to it. Each statement contains information given in one of the paragraphs. Identify the paragraph from which the information is derived.",
                                "groups": match_groups,
                            },
                            {
                                "title": "Section C 仔细阅读",
                                "type": "仔细阅读",
                                "qid_range": "47-56",
                                "instruction": "Directions: There are 2 passages in this section. Each passage is followed by some questions or unfinished statements. For each of them there are four choices marked A), B), C) and D). You should decide on the best choice.",
                                "groups": reading_groups,
                            },
                        ],
                    },
                    {
                        "part": "Part IV 翻译",
                        "type": "翻译",
                        "duration": "30分钟",
                        "instruction": "Directions: For this part, you are allowed 30 minutes to translate a passage from Chinese into English.",
                        "questions": typed_items.get("翻译", []),
                    },
                ],
            }

            # ── 6. 持久化保存试卷 ──
            paper_title = req.title or f"四级模拟试卷_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            paper_content = json.dumps(paper, ensure_ascii=False)
            new_paper = Paper(
                title=paper_title,
                content=paper_content,
                difficulty=req.difficulty,
                total_score=710,
                create_time=datetime.now(),
            )
            db.add(new_paper)
            db.flush()
            paper_id = new_paper.id
            exam_code = f"EXAM-{paper_id}"

            # ── 7. 批量持久化 ExamPaperMapping ──
            # ★ 强制按 question_id 升序写入
            result_sorted = sorted(result, key=lambda x: x[0])
            for qid, qtype, section, full_score, q in result_sorted:
                db.add(ExamPaperMapping(
                    exam_id=exam_code,
                    question_id=qid,
                    question_db_id=q.id,
                    question_type=qtype,
                    section=section,
                    full_score=full_score,
                    create_time=datetime.now(),
                ))
            db.commit()

            return {
                "code": 200,
                "msg": "试卷生成成功",
                "data": paper,
                "paper_id": paper_id,
                "exam_id": exam_code,
                "total": len(result),
            }
        finally:
            db.close()
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"code": 500, "msg": "组卷失败", "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════
# 试卷管理 CRUD
# ══════════════════════════════════════════════════════════════════════════

@router.get("/api/paper/list")
def get_paper_list():
    db = SessionLocal()
    try:
        papers = db.query(Paper).order_by(Paper.create_time.desc()).all()
        result = [{
            "id": p.id, "title": p.title, "difficulty": p.difficulty,
            "total_score": p.total_score,
            "create_time": p.create_time.strftime('%Y-%m-%d %H:%M:%S'),
        } for p in papers]
        return {"code": 200, "msg": "success", "data": result}
    finally:
        db.close()


@router.get("/api/paper/{paper_id}")
def get_paper(paper_id: int):
    db = SessionLocal()
    try:
        paper = db.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            return {"code": 404, "msg": "试卷不存在"}
        content = json.loads(paper.content) if paper.content else {}
        return {
            "code": 200, "msg": "success",
            "data": {
                "id": paper.id, "title": paper.title, "content": content,
                "difficulty": paper.difficulty, "total_score": paper.total_score,
                "create_time": paper.create_time.strftime('%Y-%m-%d %H:%M:%S'),
            },
        }
    finally:
        db.close()


@router.delete("/api/paper/{paper_id}")
def delete_paper(paper_id: int, teacher = Depends(require_teacher)):
    db = SessionLocal()
    try:
        paper = db.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            return {"code": 404, "msg": "试卷不存在"}
        # 同时删除关联映射
        exam_code = f"EXAM-{paper_id}"
        db.query(ExamPaperMapping).filter(ExamPaperMapping.exam_id == exam_code).delete()
        db.delete(paper)
        db.commit()
        return {"code": 200, "msg": "删除成功"}
    except Exception as e:
        db.rollback()
        return {"code": 500, "msg": "删除失败", "error": str(e)}
    finally:
        db.close()


# ══════════════════════════════════════════════════════════════════════════
# 在线考试：从已组试卷加载 → 学生答题 → 提交
# ══════════════════════════════════════════════════════════════════════════

@router.get("/api/exam/paper/{paper_id}")
def start_exam_from_paper(paper_id: int):
    """从已组好的试卷加载考试数据。学生端通过此接口获取预组试卷进行答题。
    通过 exam_paper_mapping 表获取稳定映射关系。
    """
    db = SessionLocal()
    try:
        paper = db.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            return {"code": 404, "msg": "试卷不存在"}

        exam_id = f"EXAM-{paper_id}"

        # 从映射表加载
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == exam_id
        ).order_by(ExamPaperMapping.question_id).all()

        if not mappings:
            # 无映射则从 paper content 重建
            content = json.loads(paper.content) if paper.content else {}
            return {
                "code": 200,
                "exam_id": exam_id,
                "title": paper.title,
                "difficulty": paper.difficulty,
                "paper": content,
                "total": 0,
                "questions_flat": [],
                "msg": "试卷已加载，但映射表为空（旧版试卷）。请重新组卷。",
            }

        questions_flat = []
        for m in mappings:
            q = db.query(DbQuestion).filter(DbQuestion.id == m.question_db_id).first()
            if q:
                content, options = _reconstruct_question(q)
                questions_flat.append({
                    "qid": m.question_id,
                    "db_id": m.question_db_id,
                    "type": m.question_type,
                    "section": m.section,
                    "content": content,
                    "options": options,
                    "answer": q.answer or "",
                    "full_score": m.full_score,
                })

        # 同时返回完整的 paper content（含 passage groups）
        paper_content = json.loads(paper.content) if paper.content else {}

        return {
            "code": 200,
            "exam_id": exam_id,
            "title": paper.title,
            "difficulty": paper.difficulty,
            "paper": paper_content,
            "total": len(questions_flat),
            "questions_flat": questions_flat,
        }
    except Exception as e:
        return {"code": 500, "msg": "加载失败", "error": str(e)}
    finally:
        db.close()


class AnswerItem(BaseModel):
    question_id: int
    answer_text: str = ""


class ExamSubmitReq(BaseModel):
    exam_id: str
    student_name: str
    answers: list[AnswerItem]


@router.post("/api/exam/submit")
def submit_exam(req: ExamSubmitReq):
    """提交答案 → 存入 student_answer 表。
    使用 exam_paper_mapping 表验证 question_id 合法性。
    """
    db = SessionLocal()
    try:
        # 验证考试存在
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == req.exam_id
        ).all()
        if not mappings:
            return {"code": 404, "msg": "考试不存在或已过期"}

        valid_qids = {m.question_id for m in mappings}
        saved = 0
        for ans in req.answers:
            if ans.question_id not in valid_qids:
                continue
            db.add(StudentAnswer(
                exam_id=req.exam_id,
                student_name=req.student_name,
                question_id=ans.question_id,
                answer_text=ans.answer_text,
            ))
            saved += 1
        db.commit()
        return {"code": 200, "msg": f"提交成功，共 {saved} 题", "count": saved}
    except Exception as e:
        db.rollback()
        return {"code": 500, "msg": "提交失败", "error": str(e)}
    finally:
        db.close()


@router.get("/api/exam/result")
def get_exam_result(exam_id: str, student_name: str):
    """查看答题结果，关联映射表获取原题信息"""
    db = SessionLocal()
    try:
        rows = db.query(StudentAnswer).filter(
            StudentAnswer.exam_id == exam_id,
            StudentAnswer.student_name == student_name,
        ).order_by(StudentAnswer.question_id).all()

        if not rows:
            return {"code": 404, "msg": "未找到答题记录"}

        # 获取映射表
        mapping_map = {}
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == exam_id
        ).all()
        for m in mappings:
            mapping_map[m.question_id] = m

        details = []
        for r in rows:
            m = mapping_map.get(r.question_id)
            item = {
                "question_id": r.question_id,
                "question_db_id": m.question_db_id if m else None,
                "student_answer": r.answer_text,
                "question_type": m.question_type if m else "",
                "full_score": m.full_score if m else 0,
            }
            # 尝试获取正确答案
            if m:
                q = db.query(DbQuestion).filter(DbQuestion.id == m.question_db_id).first()
                if q:
                    item["correct_answer"] = q.answer or ""
            details.append(item)

        return {
            "code": 200,
            "data": {
                "exam_id": exam_id,
                "student_name": student_name,
                "total_answered": len(details),
                "details": details,
            },
        }
    except Exception as e:
        return {"code": 500, "error": str(e)}
    finally:
        db.close()


# ── 在线考试：获取可用试卷列表 ──────────────────────────────────

@router.get("/api/exam/paper-list")
def get_exam_paper_list():
    """列出 exam_paper_mapping 中可用的考试编号"""
    db = SessionLocal()
    try:
        rows = db.query(ExamPaperMapping.exam_id).distinct().order_by(ExamPaperMapping.exam_id).all()
        exam_ids = [r[0] for r in rows if r[0]]
        # 获取每个exam_id的题数和满分
        result = []
        for eid in exam_ids:
            count = db.query(ExamPaperMapping).filter(ExamPaperMapping.exam_id == eid).count()
            total_score = db.query(func.sum(ExamPaperMapping.full_score)).filter(
                ExamPaperMapping.exam_id == eid).scalar() or 0
            # 取第一个mapping的create_time作为试卷时间
            first = db.query(ExamPaperMapping).filter(ExamPaperMapping.exam_id == eid).first()
            result.append({
                "exam_id": eid,
                "question_count": count,
                "total_score": round(total_score, 1),
            })
        return {"code": 200, "data": result}
    except Exception as e:
        return {"code": 500, "error": str(e)}
    finally:
        db.close()


# ── 在线考试：开始考试（加载试卷题目）─────────────────────────

class ExamStartReq(BaseModel):
    exam_id: str
    student_name: str
    class_name: str = ""


@router.post("/api/exam/start")
def start_exam(req: ExamStartReq):
    """根据exam_id从exam_paper_mapping加载试卷题目，返回扁平题目列表"""
    db = SessionLocal()
    try:
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == req.exam_id
        ).order_by(ExamPaperMapping.question_id).all()

        if not mappings:
            return {"code": 404, "msg": "考试编号不存在或已过期"}

        questions = []
        for m in mappings:
            q = db.query(DbQuestion).filter(DbQuestion.id == m.question_db_id).first()
            if not q:
                continue
            content, options = _reconstruct_question(q)
            qtype = m.question_type
            # 判断是否为主观题（写作/翻译）
            is_subjective = qtype in ("写作", "翻译")
            questions.append({
                "qid": m.question_id,
                "db_id": m.question_db_id,
                "type": qtype,
                "section": m.section,
                "content": content,
                "options": options,
                "answer": q.answer or "",
                "full_score": m.full_score,
                "is_subjective": is_subjective,
            })

        return {
            "code": 200,
            "msg": "试卷加载成功",
            "exam_id": req.exam_id,
            "student_name": req.student_name,
            "class_name": req.class_name,
            "questions": questions,
            "total": len(questions),
            "total_score": sum(q["full_score"] for q in questions),
        }
    except Exception as e:
        return {"code": 500, "msg": "加载试卷失败", "error": str(e)}
    finally:
        db.close()


# ── AI 评分（主观题：写作、翻译）─────────────────────────────

def _grade_subjective_by_ai(question_text, student_answer, full_score, qtype):
    """使用 DeepSeek 对主观题评分，返回 0 ~ full_score 的分数"""
    try:
        from openai import OpenAI
        import os
        from dotenv import load_dotenv
        load_dotenv()
        client = OpenAI(
            api_key=os.getenv("DEEPSEEK_API_KEY"),
            base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
        )
        model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

        if qtype == "写作":
            system_prompt = """你是一位大学英语四级（CET-4）作文阅卷专家。请严格按照CET-4评分标准对学生作文进行评分。
评分维度（满分106.5分）：
1. 内容完整性 (0-35分)：是否覆盖题目要求的所有要点
2. 语言准确性 (0-35分)：语法、拼写、标点是否正确
3. 词汇与句式 (0-36.5分)：词汇丰富度、句式多样性

评分规则：
- 严格按上述维度打分，最后给出总分
- 120-180词为合理长度，过短或过长适当扣分
- 只返回JSON格式，不要其他文字"""
            user_prompt = f"""请对以下CET-4作文进行评分。

题目：
{question_text}

学生作文：
{student_answer}

满分：{full_score}分

返回JSON格式：
{{"content_score": 数值, "language_score": 数值, "vocabulary_score": 数值, "total_score": 数值, "comment": "简评"}}"""
        else:  # 翻译
            system_prompt = """你是一位大学英语四级（CET-4）翻译阅卷专家。请严格按照CET-4评分标准对学生翻译进行评分。
评分维度（满分106.5分）：
1. 准确性 (0-40分)：译文是否准确传达原文意思
2. 流畅度 (0-35分)：译文是否通顺、符合英语表达习惯
3. 语法词汇 (0-31.5分)：语法是否正确、词汇使用是否恰当

评分规则：
- 严格按上述维度打分，最后给出总分
- 只返回JSON格式，不要其他文字"""
            user_prompt = f"""请对以下CET-4翻译进行评分。

需要翻译的中文原文：
{question_text}

学生的英文翻译：
{student_answer}

满分：{full_score}分

返回JSON格式：
{{"accuracy_score": 数值, "fluency_score": 数值, "grammar_score": 数值, "total_score": 数值, "comment": "简评"}}"""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=500,
            response_format={"type": "json_object"},
        )
        raw = response.choices[0].message.content
        result = json.loads(raw)
        total = float(result.get("total_score", 0))
        # 确保不超过满分
        return min(total, full_score)
    except Exception:
        # AI评分失败时返回0
        return 0.0


# ── 在线考试：提交并评分 ──────────────────────────────────────

class ExamGradeReq(BaseModel):
    exam_id: str
    student_name: str
    class_name: str = ""
    answers: list[AnswerItem]


@router.post("/api/exam/grade")
def grade_exam(req: ExamGradeReq):
    """
    提交答案 + 自动评分：
    1. 保存答案到 student_answer 表（若已存在则跳过）
    2. 客观题（听力/选词填空/段落匹配/仔细阅读）：直接匹配答案，正确给满分，否则0分
    3. 主观题（写作/翻译）：调用AI评分
    4. 保存评分结果到 exam_record 表
    """
    db = SessionLocal()
    try:
        # 获取映射表
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == req.exam_id
        ).all()
        if not mappings:
            return {"code": 404, "msg": "考试编号不存在"}

        mapping_map = {m.question_id: m for m in mappings}
        valid_qids = set(mapping_map.keys())

        # 从题库加载题目（用于获取正确答案）
        db_ids = set(m.question_db_id for m in mappings)
        questions_map = {}
        for db_id in db_ids:
            q = db.query(DbQuestion).filter(DbQuestion.id == db_id).first()
            if q:
                questions_map[db_id] = q

        # 构建答案映射（方便快速查找）
        answer_map = {a.question_id: a.answer_text for a in req.answers}

        results = []
        total_score = 0
        total_full = 0
        graded_count = 0

        # 遍历试卷所有题目（包括未作答的）
        for m in mappings:
            qid = m.question_id
            q = questions_map.get(m.question_db_id)
            qtype = m.question_type or ""
            full_score = m.full_score or 1
            is_subjective = qtype in ("写作", "翻译")

            answer_text = answer_map.get(qid, "") or ""

            # 保存答案到 student_answer（如已存在则跳过）
            if answer_text.strip():
                existing = db.query(StudentAnswer).filter(
                    StudentAnswer.exam_id == req.exam_id,
                    StudentAnswer.student_name == req.student_name,
                    StudentAnswer.question_id == qid,
                ).first()
                if not existing:
                    db.add(StudentAnswer(
                        exam_id=req.exam_id,
                        student_name=req.student_name,
                        question_id=qid,
                        answer_text=answer_text,
                    ))

            # 评分
            score = 0

            if is_subjective:
                # 主观题：AI评分
                if answer_text.strip():
                    question_content = q.content if q else ""
                    score = _grade_subjective_by_ai(question_content, answer_text, full_score, qtype)
                    score = round(score, 1)
                else:
                    score = 0  # 未作答
            else:
                # 客观题：直接匹配答案
                if answer_text.strip():
                    correct = (q.answer or "").strip().upper()
                    student_ans = answer_text.strip().upper()
                    if correct and student_ans == correct:
                        score = full_score
                    else:
                        score = 0
                else:
                    score = 0

            # 保存到 exam_record（如已存在则更新）
            exist_record = db.query(ExamRecord).filter(
                ExamRecord.exam_id == req.exam_id,
                ExamRecord.student_name == req.student_name,
                ExamRecord.question_id == qid,
            ).first()
            if exist_record:
                exist_record.score = score
                exist_record.full_score = full_score
                exist_record.class_name = req.class_name
            else:
                db.add(ExamRecord(
                    exam_id=req.exam_id,
                    student_name=req.student_name,
                    class_name=req.class_name,
                    question_id=qid,
                    score=score,
                    full_score=full_score,
                ))

            total_score += score
            total_full += full_score
            graded_count += 1

            results.append({
                "question_id": qid,
                "type": qtype,
                "full_score": full_score,
                "score": score,
                "is_subjective": is_subjective,
                "correct_answer": q.answer if q and not is_subjective else None,
                "student_answer": answer_text,
            })

        db.commit()

        total_full = round(total_full, 1)
        rate = round(total_score / total_full * 100, 1) if total_full else 0
        return {
            "code": 200,
            "msg": f"评分完成，共{graded_count}题，总分{round(total_score, 1)}/{total_full}（得分率{rate}%）",
            "data": {
                "exam_id": req.exam_id,
                "student_name": req.student_name,
                "total_score": round(total_score, 1),
                "total_full": total_full,
                "score_rate": rate,
                "graded_count": graded_count,
                "details": results,
            },
        }
    except Exception as e:
        db.rollback()
        import traceback
        traceback.print_exc()
        return {"code": 500, "msg": "评分失败", "error": str(e)}
    finally:
        db.close()


# ── 查看学生答卷（student_answer 原始答案）─────────────────────

@router.get("/api/exam/student-answers")
def get_student_answers(exam_id: str, student_name: str):
    """查询指定学生在某次考试中的原始答案和批改结果"""
    db = SessionLocal()
    try:
        # 获取该学生的答题记录
        answers = db.query(StudentAnswer).filter(
            StudentAnswer.exam_id == exam_id,
            StudentAnswer.student_name == student_name,
        ).order_by(StudentAnswer.question_id).all()

        if not answers:
            return {"code": 404, "msg": "该学生没有答题记录（可能为Excel导入的模拟数据）"}

        # 获取映射表
        mappings = db.query(ExamPaperMapping).filter(
            ExamPaperMapping.exam_id == exam_id
        ).all()
        mapping_map = {m.question_id: m for m in mappings}

        # 获取评分结果
        records = db.query(ExamRecord).filter(
            ExamRecord.exam_id == exam_id,
            ExamRecord.student_name == student_name,
        ).all()
        score_map = {r.question_id: r for r in records}

        # 组装数据
        details = []
        for ans in answers:
            qid = ans.question_id
            m = mapping_map.get(qid)
            sc = score_map.get(qid)
            qtype = m.question_type if m else ""
            is_subjective = qtype in ("写作", "翻译")

            # 获取正确答案和题目内容
            correct_answer = None
            question_content = ""
            if m:
                q = db.query(DbQuestion).filter(DbQuestion.id == m.question_db_id).first()
                if q:
                    if not is_subjective:
                        correct_answer = q.answer or ""
                    question_content, _ = _reconstruct_question(q)

            details.append({
                "question_id": qid,
                "type": qtype,
                "question_content": question_content,
                "student_answer": ans.answer_text or "",
                "correct_answer": correct_answer,
                "score": sc.score if sc else 0,
                "full_score": m.full_score if m else (sc.full_score if sc else 0),
                "is_subjective": is_subjective,
                "submitted_at": ans.submitted_at.strftime("%Y-%m-%d %H:%M:%S") if ans.submitted_at else "",
            })

        # 统计
        total_score = sum(d["score"] for d in details)
        total_full = sum(d["full_score"] for d in details)
        total_full = round(total_full, 1)
        rate = round(total_score / total_full * 100, 1) if total_full else 0

        return {
            "code": 200,
            "data": {
                "exam_id": exam_id,
                "student_name": student_name,
                "total_score": round(total_score, 1),
                "total_full": total_full,
                "score_rate": rate,
                "question_count": len(details),
                "details": details,
            },
        }
    except Exception as e:
        return {"code": 500, "error": str(e)}
    finally:
        db.close()


# ══════════════════════════════════════════════════════════════════════════
# Word 导出
# ══════════════════════════════════════════════════════════════════════════

os.makedirs("temp", exist_ok=True)


@router.post("/api/paper/export")
def export_paper(req: PaperGenerateReq, teacher = Depends(require_teacher)):
    """导出 Word 试卷（独立抽题，不保存到数据库）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    db = SessionLocal()
    try:
        all_qs = db.query(DbQuestion).all()

        pools = {}
        for display_type in ["写作", "听力", "选词填空", "段落匹配", "仔细阅读", "翻译"]:
            pools[display_type] = _find_pool(all_qs, display_type) or all_qs

        type_needed = {}
        for _, qtype, _, _ in CET4_LAYOUT:
            type_needed[qtype] = type_needed.get(qtype, 0) + 1

        type_selected = {}
        for qtype, need in type_needed.items():
            pool = pools.get(qtype, all_qs)
            selected = _pick_by_ratio(pool, need, req.difficulty)
            selected.sort(key=lambda q: (getattr(q, 'question_number', 0) or 0, q.id))
            type_selected[qtype] = selected

        # 按 CET4 顺序编排
        type_idx = {}
        export_data = OrderedDict()
        for qid, qtype, section, full_score in CET4_LAYOUT:
            lst = type_selected.get(qtype, [])
            idx = type_idx.get(qtype, 0)
            if idx < len(lst):
                q = lst[idx]
                type_idx[qtype] = idx + 1
                export_data.setdefault(qtype, []).append({
                    "id": q.id, "type": qtype, "content": q.content,
                    "answer": q.answer, "analysis": q.analysis,
                    "options": q.options,
                })

        doc = Document()
        title = doc.add_paragraph()
        title_run = title.add_run("大学英语四级模拟试卷")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"难度等级：{req.difficulty}级  |  总分：710分")
        sub_run.font.size = Pt(12)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        for type_name, questions in export_data.items():
            if not questions:
                continue
            type_para = doc.add_paragraph()
            type_run = type_para.add_run(f"{type_name}（共{len(questions)}题）")
            type_run.bold = True
            type_run.font.size = Pt(14)
            type_run.font.color.rgb = RGBColor(0, 0, 128)

            for idx_q, q in enumerate(questions, 1):
                p = doc.add_paragraph()
                p_run = p.add_run(f"{idx_q}. {q['content']}")
                p_run.font.size = Pt(11)
                if q.get("options"):
                    try:
                        opts = json.loads(q["options"]) if isinstance(q["options"], str) else q["options"]
                    except Exception:
                        opts = q["options"]
                    if isinstance(opts, list):
                        for opt in opts:
                            if isinstance(opt, dict):
                                opt_text = f"{opt.get('label', '')}) {opt.get('text', '')}"
                            else:
                                opt_text = str(opt)
                            opt_p = doc.add_paragraph()
                            opt_run = opt_p.add_run(f"    {opt_text}")
                            opt_run.font.size = Pt(11)
                doc.add_paragraph("")

        filename = f"四级模拟试卷_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = os.path.join("temp", filename)
        doc.save(file_path)

        return FileResponse(
            file_path, filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"code": 500, "msg": "导出失败", "error": str(e)}
    finally:
        db.close()
